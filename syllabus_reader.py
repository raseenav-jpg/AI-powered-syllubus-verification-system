import re
import os
import pdfplumber

from docx import Document


# ---------------------------------------------------------
# TEXT EXTRACTION
# ---------------------------------------------------------

def extract_text_from_docx(file_path):

    doc = Document(file_path)

    parts = []

    # Paragraphs
    for para in doc.paragraphs:

        if para.text.strip():

            parts.append(
                para.text.strip()
            )

    # Tables
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                text = cell.text.strip()

                if text:

                    parts.append(text)

    return "\n".join(parts)


def extract_text_from_pdf(file_path):

    parts = []

    with pdfplumber.open(file_path) as pdf:

        for page in pdf.pages:

            txt = page.extract_text()

            if txt:

                parts.append(txt)

    return "\n".join(parts)


def extract_text(file_path):

    extension = os.path.splitext(
        file_path
    )[1].lower()

    # DOCX
    if extension == ".docx":

        return extract_text_from_docx(
            file_path
        )

    # PDF
    elif extension == ".pdf":

        return extract_text_from_pdf(
            file_path
        )

    else:

        raise ValueError(
            f"Unsupported file type: {extension}"
        )


# ---------------------------------------------------------
# TABLE EXTRACTION
# ---------------------------------------------------------

def extract_course_details_from_tables(file_path):

    details = {

        "credits": None,
        "hours_per_week": None
    }

    extension = os.path.splitext(
        file_path
    )[1].lower()

    # Only DOCX supports table parsing
    if extension != ".docx":

        return details

    try:

        doc = Document(file_path)

        for table in doc.tables:

            rows = []

            for row in table.rows:

                row_data = []

                for cell in row.cells:

                    value = cell.text.strip()

                    row_data.append(value)

                rows.append(row_data)

            # Search for Course Details table
            for i, row in enumerate(rows):

                row_text = " ".join(row).lower()

                if (
                    "credit" in row_text
                    and "lecture per week" in row_text
                ):

                    # Next row contains values
                    if i + 1 < len(rows):

                        values = rows[i + 1]

                        nums = []

                        for v in values:

                            found = re.findall(
                                r'\d+',
                                v
                            )

                            nums.extend(found)

                        print(
                            "\nTABLE NUMBERS:",
                            nums
                        )

                        # Expected:
                        # [4, 3, 2, 75]

                        if len(nums) >= 4:

                            credits = int(nums[0])

                            lecture = int(nums[1])

                            practical = int(nums[2])

                            details["credits"] = credits

                            details["hours_per_week"] = (
                                lecture + practical
                            )

                            return details

    except Exception as e:

        print(
            "TABLE EXTRACTION ERROR:",
            e
        )

    return details


# ---------------------------------------------------------
# NORMALIZATION
# ---------------------------------------------------------

def normalize_text(text):

    text = text.lower()

    text = re.sub(
        r'\s+',
        ' ',
        text
    )

    return text.strip()


# ---------------------------------------------------------
# SECTION DETECTION
# ---------------------------------------------------------

def has_any_section(text, possible_names):

    normalized_text = normalize_text(text)

    for name in possible_names:

        if normalize_text(name) in normalized_text:

            return True

    return False


# ---------------------------------------------------------
# COURSE OUTCOMES
# ---------------------------------------------------------

def count_course_outcomes(text):

    matches = re.findall(
        r'\bCO\s*[1-8]\b',
        text,
        re.IGNORECASE
    )

    unique_cos = set(

        m.upper().replace(" ", "")

        for m in matches
    )

    return len(unique_cos)


# ---------------------------------------------------------
# BASIC COURSE DETAILS
# ---------------------------------------------------------

def extract_basic_course_details(
    text,
    file_path
):

    details = {

        "discipline": None,
        "course_code": None,
        "course_title": None,
        "type_of_course": None,
        "credits": None,
        "hours_per_week": None
    }

    clean_text = " ".join(
        text.split()
    )

    # -------------------------------------------------
    # COURSE CODE
    # -------------------------------------------------

    match = re.search(

        r'Course\s*Code\s*([A-Z0-9]+)',

        clean_text,

        re.IGNORECASE
    )

    if match:

        details["course_code"] = (
            match.group(1)
        )

   
    # -------------------------------------------------
    # COURSE TITLE
    # -------------------------------------------------

    title_match = re.search(

        r'Course\s*Code\s*[A-Z0-9]+.*?'
        r'Course\s*Title\s*(.*?)'
        r'Type\s*of\s*Course',

        clean_text,

        re.IGNORECASE
    )

    if title_match:

        extracted_title = (
            title_match.group(1)
            .strip()
        )

        # Remove extra spaces
        extracted_title = re.sub(
            r'\s+',
            ' ',
            extracted_title
        )

        details["course_title"] = (
            extracted_title
        )
    # -------------------------------------------------
    # TYPE OF COURSE
    # -------------------------------------------------

    match = re.search(

        r'Type\s*of\s*Course\s*([A-Za-z ]+)',

        clean_text,

        re.IGNORECASE
    )

    if match:

        extracted = (
            match.group(1)
            .strip()
            .split()[0]
        )

        details["type_of_course"] = (
            extracted.upper()
        )

        # -------------------------------------------------
    # Credits and Hours
    # Will be loaded automatically from rules.py
    # after course classification
    # -------------------------------------------------

    details["credits"] = None
    details["hours_per_week"] = None
    return details

# ---------------------------------------------------------
# MAIN EXTRACTION FUNCTION
# ---------------------------------------------------------

def extract_syllabus_info(file_path):

    text = extract_text(file_path)

    basic_details = (
        extract_basic_course_details(
            text,
            file_path
        )
    )

    info = {

        "text": text,

        # Basic Details
        "discipline":
        basic_details["discipline"],

        "course_code":
        basic_details["course_code"],

        "course_title":
        basic_details["course_title"],

        "type_of_course":
        basic_details["type_of_course"],

        "credits":
        basic_details["credits"],

        "hours_per_week":
        basic_details["hours_per_week"],

        # Sections
        "has_course_summary":
        has_any_section(
            text,
            ["Course Summary"]
        ),

        "has_course_outcomes":
        has_any_section(
            text,
            [
                "Course Outcomes",
                "Course Outcomes (CO)"
            ]
        ),

        "has_detailed_syllabus":
        (
            has_any_section(
                text,
                ["Detailed Syllabus"]
            )
            or
            (
                has_any_section(
                    text,
                    ["Module"]
                )
                and
                has_any_section(
                    text,
                    ["Unit"]
                )
            )
        ),

        "has_references":
        has_any_section(
            text,
            [
                "Reference",
                "References"
            ]
        ),

        # CO Count
        "co_count":
        count_course_outcomes(text)
    }

    print("\n========== FINAL INFO ==========")
    print(info)
    print("================================\n")

    return info
def extract_syllabus_info_from_text(text):

    clean_text = " ".join(
        text.split()
    )

    info = {}

    # -------------------------------------------------
    # COURSE CODE
    # -------------------------------------------------

    match = re.search(

        r'Course\s*Code\s*([A-Z0-9]+)',

        clean_text,

        re.IGNORECASE
    )

    info["course_code"] = (
        match.group(1)
        if match else None
    )

    # -------------------------------------------------
    # COURSE TITLE
    # -------------------------------------------------

    title_match = re.search(

        r'Course\s*Code\s*[A-Z0-9]+.*?'
        r'Course\s*Title\s*(.*?)'
        r'Type\s*of\s*Course',

        clean_text,

        re.IGNORECASE
    )

    if title_match:

        extracted_title = (
            title_match.group(1)
            .strip()
        )

        extracted_title = re.sub(
            r'\s+',
            ' ',
            extracted_title
        )

        info["course_title"] = (
            extracted_title
        )

    else:

        info["course_title"] = None

    # -------------------------------------------------
    # TYPE OF COURSE
    # -------------------------------------------------

    match = re.search(

        r'Type\s*of\s*Course\s*([A-Za-z ]+)',

        clean_text,

        re.IGNORECASE
    )

    info["type_of_course"] = (

        match.group(1)
        .strip()
        .split()[0]
        .upper()

        if match else None
    )

    # -------------------------------------------------
    # SECTIONS
    # -------------------------------------------------

    lower = clean_text.lower()

    info["has_course_summary"] = (
        "course summary" in lower
    )

    info["has_course_outcomes"] = (
        "course outcomes" in lower
    )

    info["has_detailed_syllabus"] = (
        "detailed syllabus" in lower
        or "module" in lower
    )

    info["has_references"] = (
        "reference" in lower
    )

    # -------------------------------------------------
    # COURSE OUTCOMES
    # -------------------------------------------------

    cos = re.findall(
        r'CO\s*[1-8]',
        clean_text,
        re.IGNORECASE
    )

    info["co_count"] = len(set(cos))

    return info
