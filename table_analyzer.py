from docx import Document


def extract_tables(file_path):
    """
    Extract all tables from a DOCX file.
    Returns a list of dictionaries containing:
    - table number
    - row count
    - column count
    - cell data
    """

    doc = Document(file_path)
    all_tables = []

    for table_index, table in enumerate(doc.tables, start=1):
        table_data = {
            "table_number": table_index,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "data": []
        }

        for row in table.rows:
            row_data = []

            for cell in row.cells:
                cell_text = "\n".join(
                    p.text.strip()
                    for p in cell.paragraphs
                    if p.text.strip()
                )

                row_data.append(cell_text)

            table_data["data"].append(row_data)

        all_tables.append(table_data)

    return all_tables
