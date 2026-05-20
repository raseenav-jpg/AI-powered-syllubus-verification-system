from docx import Document


def get_first_run_format(paragraph):
    """
    Extract formatting from the first non-empty run in a paragraph.
    """
    for run in paragraph.runs:
        if run.text.strip():
            return {
                "font_name": run.font.name,
                "font_size": (
                    run.font.size.pt if run.font.size else None
                ),
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "alignment": (
                    str(paragraph.alignment)
                    if paragraph.alignment is not None
                    else "None"
                )
            }

    return None


def extract_all_formats(file_path):
    """
    Extract formatting from all paragraphs and all table cell paragraphs.
    """
    doc = Document(file_path)
    formats = []

    # Normal paragraphs
    for para in doc.paragraphs:
        fmt = get_first_run_format(para)
        if fmt:
            formats.append(fmt)

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fmt = get_first_run_format(para)
                    if fmt:
                        formats.append(fmt)

    return formats


def compare_formatting(template_path, target_path):
    """
    Compare formatting sequentially between template and target.
    """
    template_formats = extract_all_formats(template_path)
    target_formats = extract_all_formats(target_path)

    results = []
    count = min(len(template_formats), len(target_formats))

    for i in range(count):
        t = template_formats[i]
        u = target_formats[i]

        for key in t:
            if t[key] != u[key]:
                results.append(
                    f"✗ Format difference at item {i+1}: "
                    f"{key} "
                    f"(Template: {t[key]}, Target: {u[key]})"
                )

    if not results:
        results.append("✓ Formatting matches")

    return results
